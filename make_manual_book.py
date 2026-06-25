import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_manual():
    doc = docx.Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # charcoal/dark gray

    # Helpers
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Navy Blue
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # Slate Blue
        p.paragraph_format.space_after = Pt(30)

    def add_divider():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("____________________________________________________")
        run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1) # light gray border
        p.paragraph_format.space_after = Pt(24)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Deep Navy

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) # Teal

    def add_body(text, bold_prefix=None, indent=0, space_after=6, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if indent > 0:
            p.paragraph_format.left_indent = Inches(indent * 0.25)
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(11)
            r_bold.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.italic = italic

    def add_bullet(text, bold_prefix=None, indent=1):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(11)
            r_bold.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_number(text, bold_prefix=None, indent=1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(indent * 0.25)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(11)
            r_bold.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_note_box(title, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        
        run_title = p.add_run(title + "\n")
        run_title.bold = True
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(10.5)
        run_title.font.color.rgb = RGBColor(0xB4, 0x53, 0x09) # Amber/Gold
        
        run_text = p.add_run(text)
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(10.5)
        run_text.font.italic = True
        run_text.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # --- COVER / TITLE ---
    add_title("Buku Panduan Penggunaan Aplikasi Absenku")
    add_subtitle("Sistem Absensi Online Sekolah Berbasis Lokasi (GPS)")
    add_divider()

    # --- SECT 1: PENDAHULUAN ---
    add_h1("1. Pendahuluan")
    add_body("Absenku adalah aplikasi absensi sekolah online yang modern, cepat, dan mudah digunakan. Aplikasi ini dirancang agar proses pencatatan kehadiran di sekolah berjalan lebih praktis dan akurat menggunakan bantuan GPS/Lokasi HP masing-masing.")
    add_body("Buku panduan ini disusun dengan bahasa yang sederhana untuk membantu Siswa, Guru, Wali Kelas, Petugas Piket, dan Admin Sekolah dalam menggunakan seluruh fitur yang tersedia. Semoga aplikasi ini mempermudah kegiatan belajar mengajar kita!")

    # --- SECT 2: PERSIAPAN AWAL ---
    add_h1("2. Persiapan Penting Sebelum Menggunakan Aplikasi")
    add_body("Agar aplikasi dapat mendeteksi kehadiran Anda dengan sukses, pastikan HP Anda memenuhi persyaratan berikut:")
    add_bullet("Koneksi Internet: HP Anda harus terhubung dengan paket data internet aktif atau Wi-Fi sekolah.", "1. ")
    add_bullet("GPS / Lokasi Aktif: Pastikan fitur 'Lokasi' atau 'GPS' di HP Anda sudah dihidupkan.", "2. ")
    add_bullet("Izin Lokasi di Browser: Saat membuka aplikasi absensi untuk pertama kali, HP Anda akan memunculkan pertanyaan seperti 'Apakah Anda mengizinkan situs ini mengakses lokasi Anda?'. Anda WAJIB memilih IZINKAN (ALLOW). Jika tidak, sistem tidak akan bisa mencatat absen Anda.", "3. ")

    # --- SECT 3: PANDUAN SISWA ---
    add_h1("3. Panduan untuk Siswa")
    add_body("Sebagai siswa, tugas utama Anda adalah mencatat kehadiran harian secara mandiri di sekolah dan mengajukan izin jika berhalangan hadir.")

    add_h2("A. Cara Login Pertama Kali & Pengaturan Profil Wajib")
    add_number("Buka browser (Google Chrome untuk Android atau Safari untuk iPhone) lalu masuk ke alamat website Absenku yang diberikan pihak sekolah.", "1. ")
    add_number("Masukkan Email/NISN Anda dan Password default yang diberikan sekolah, lalu klik tombol Login.", "2. ")
    add_number("Bila Anda baru pertama kali login, sistem secara otomatis akan mengarahkan Anda ke halaman edit profil. Di halaman ini, Anda diwajibkan untuk mengubah password default Anda serta mengisi nomor HP (WhatsApp) aktif Anda.", "3. ")
    add_number("Klik tombol Simpan untuk memperbarui informasi profil dan mengaktifkan akun Anda agar dapat mengakses halaman absensi.", "4. ")

    add_h2("B. Cara Absen Masuk (Saat Datang di Sekolah)")
    add_number("Pastikan Anda sudah berada di dalam lingkungan sekolah (masuk dalam radius wilayah sekolah yang ditentukan, misalnya maksimal 50 meter dari titik tengah sekolah).", "1. ")
    add_number("Buka halaman absensi utama di HP Anda, lalu ketuk tombol Absen Masuk.", "2. ")
    add_number("Tunggu beberapa detik saat sistem mendeteksi titik koordinat lokasi GPS Anda.", "3. ")
    add_number("Jika lokasi Anda cocok dengan wilayah sekolah, sistem akan mencatat waktu kedatangan Anda dan menampilkan status kehadiran Anda.", "4. ")
    add_note_box("Penting untuk Siswa Terlambat:", "Batas waktu absen masuk disesuaikan dengan aturan sekolah. Jika Anda menekan absen masuk lewat dari jam masuk sekolah ditambah masa toleransi yang ditentukan, sistem akan mencatat status kehadiran Anda secara otomatis sebagai 'Terlambat'.")

    add_h2("C. Cara Absen Pulang (Saat Selesai Sekolah)")
    add_number("Saat jam pulang sekolah tiba, buka kembali website Absenku di HP Anda.", "1. ")
    add_number("Ketuk tombol Absen Pulang di halaman utama.", "2. ")
    add_number("Sistem akan kembali memverifikasi lokasi GPS Anda (pastikan Anda masih berada di lingkungan sekolah sebelum pulang).", "3. ")
    add_number("Waktu pulang Anda akan tercatat secara otomatis.", "4. ")

    add_h2("D. Cara Mengajukan Izin atau Sakit")
    add_body("Jika Anda terpaksa tidak bisa hadir karena sakit atau urusan mendesak keluarga, Anda harus meminta izin melalui aplikasi:")
    add_number("Gulir halaman absensi ke bagian bawah pada kartu formulir Pengajuan Izin.", "1. ")
    add_number("Pilih Jenis Izin: 'Izin Tidak Masuk' (jika tidak masuk sekolah dari pagi) atau 'Izin Pulang Lebih Awal' (jika Anda terpaksa pulang di tengah jam pelajaran).", "2. ")
    add_number("Pilih Waktu Izin: 'Hari Ini' atau 'Besok' (jika Anda ingin mengajukan izin sehari sebelumnya).", "3. ")
    add_number("Pilih Alasan Izin: 'Sakit' atau 'Urusan Penting/Mendadak'.", "4. ")
    add_number("Tulis keterangan singkat (misal: 'Demam tinggi dan sedang berobat di klinik' atau 'Mengikuti acara pernikahan keluarga').", "5. ")
    add_number("Ketuk Kirim Pengajuan. Status pengajuan izin Anda akan muncul pada kartu status dengan keterangan 'Menunggu'.", "6. ")
    add_number("PENTING: Setelah mengirimkan pengajuan di website, Anda wajib segera menghubungi Wali Kelas Anda masing-masing melalui WhatsApp untuk melakukan konfirmasi serta melampirkan berkas bukti fisik (seperti foto surat izin orang tua atau surat keterangan dokter) agar pengajuan dapat divalidasi dan diproses oleh Petugas Piket.", "7. ")

    # --- SECT 4: PANDUAN GURU ---
    add_h1("4. Panduan untuk Guru dan Wali Kelas")
    add_body("Bapak/Ibu Guru dan Wali Kelas dapat memantau kehadiran siswa secara real-time dan mengunduh laporan kehadiran.")

    add_h2("A. Memantau Kehadiran Siswa (Dashboard Guru)")
    add_number("Masuk ke website menggunakan akun Guru Anda.", "1. ")
    add_number("Anda akan melihat Dashboard Utama yang memuat grafik dan data statistik kehadiran hari ini.", "2. ")
    add_number("Gunakan filter kelas untuk melihat status kehadiran detail dari kelas yang dipilih (Hadir, Terlambat, Izin, Alfa, atau Belum Absen).", "3. ")

    add_h2("B. Memeriksa Rekap Absensi (Laporan Kehadiran)")
    add_body("Gunakan menu Rekap Absensi untuk melihat riwayat kehadiran siswa. Menu ini memiliki dua tab:")
    add_bullet("Menampilkan tabel kehadiran lengkap per tanggal. Anda dapat menyaring data berdasarkan kelas, status kehadiran tertentu, dan rentang tanggal.", "Tab Rekap Absen (Detail): ")
    add_bullet("Menampilkan akumulasi total jumlah kehadiran (Hadir, Izin, Telat, Alfa) dari masing-masing siswa selama periode tanggal yang Anda pilih.", "Tab Rekap Keterangan (Ringkasan): ")

    add_h2("C. Mengunduh Laporan Kehadiran (Excel & PDF)")
    add_number("Atur filter pencarian pada menu Rekap Absensi sesuai dengan kelas dan rentang tanggal yang diinginkan.", "1. ")
    add_number("Ketuk tombol Ekspor Excel untuk mengunduh laporan dalam format file spreadsheet Excel (.xlsx).", "2. ")
    add_number("Ketuk tombol Ekspor PDF untuk mengunduh laporan dalam format dokumen PDF (.pdf) siap cetak.", "3. ")

    # --- SECT 5: PANDUAN PETUGAS PIKET ---
    add_h1("5. Panduan untuk Petugas Piket")
    add_body("Petugas Piket bertugas meninjau dan memproses persetujuan pengajuan izin siswa.")

    add_h2("A. Memeriksa Pengajuan Izin Siswa")
    add_number("Masuk menggunakan akun Petugas Piket Anda dan klik menu Persetujuan Izin.", "1. ")
    add_number("Sistem akan menampilkan daftar antrean izin siswa yang berstatus 'Menunggu'.", "2. ")
    add_number("Periksa informasi pengajuan mulai dari nama siswa, jenis izin, alasan, hingga keterangan alasan yang ditulis siswa.", "3. ")
    add_number("PENTING: Sebelum menekan tombol tindakan, Petugas Piket wajib melakukan konfirmasi terlebih dahulu ke Wali Kelas dari siswa yang bersangkutan untuk memvalidasi surat bukti yang dikirim siswa melalui WhatsApp.", "4. ")

    add_h2("B. Menyetujui atau Menolak Pengajuan")
    add_bullet("Setelah mendapatkan konfirmasi validasi dari Wali Kelas bahwa izin tersebut benar dan sah, ketuk tombol Setujui. Status absen siswa pada hari itu akan otomatis tercatat sebagai 'Izin'.", "Pilihan Setujui: ")
    add_bullet("Jika Wali Kelas mengonfirmasi bahwa izin tersebut tidak sah/palsu, atau siswa tidak mengonfirmasi ke Wali Kelas, ketuk tombol Tolak. Status pengajuan ditolak dan siswa tetap dianggap tidak hadir sebelum melakukan absensi.", "Pilihan Tolak: ")

    # --- SECT 6: PANDUAN ADMIN ---
    add_h1("6. Panduan untuk Administrator Sekolah (Admin)")
    add_body("Admin bertugas mengelola data dasar sekolah, pengaturan absensi, serta data guru dan siswa.")

    add_h2("A. Mengatur Jam Absen & Koordinat Lokasi GPS Sekolah")
    add_number("Buka menu Pengaturan Sekolah di dashboard admin.", "1. ")
    add_number("Masukkan Nama Sekolah, koordinat Latitude, koordinat Longitude, dan Radius Toleransi (dalam satuan meter). Siswa di luar radius ini tidak akan bisa melakukan absensi.", "2. ")
    add_number("Tentukan jam buka dan tutup untuk Absen Masuk, jam buka dan tutup untuk Absen Pulang, serta durasi Toleransi Terlambat (dalam menit).", "3. ")
    add_number("Klik Simpan Pengaturan. Perubahan akan langsung berlaku pada sistem absensi siswa.", "4. ")

    add_h2("B. Mengelola Data Pengguna (Siswa & Guru)")
    add_bullet("Admin dapat menambah akun secara satuan melalui menu Kelola Siswa, Kelola Guru, atau Kelola Pengguna.", "1. ")
    add_bullet("Admin dapat mengedit profil pengguna (seperti mengubah NISN, nama, kelas, maupun nomor HP) atau menghapus akun secara permanen melalui daftar tabel pengguna.", "2. ")

    add_h2("C. Cara Mengimpor Data Siswa Secara Massal (Excel)")
    add_number("Buka menu Siswa lalu ketuk tombol Impor Siswa di sudut kanan atas.", "1. ")
    add_number("Unduh file template Excel (.xlsx) yang telah disediakan sistem.", "2. ")
    add_number("Isi data siswa di kolom yang tersedia (Nama, NISN, Kelas, Jurusan, Nomor HP Orangtua, dan Nomor HP Siswa). Jangan mengubah tata letak kolom template.", "3. ")
    add_number("Unggah file Excel tersebut pada halaman impor, kemudian klik tombol Impor.", "4. ")
    add_note_box("Info Password & Email Impor:", "Siswa yang diimpor secara otomatis dibuatkan email dengan format: [NISN]@sekolah.local dan password default: siswa123. Siswa dapat login dengan akun ini dan diwajibkan mengubah password pada login pertama.")

    add_h2("D. Mengelola Kelas & Hapus Siswa Massal (Bulk Delete)")
    add_bullet("Admin dapat menambahkan nama kelas baru beserta jurusan atau menghapusnya pada menu Kelas.", "1. ")
    add_bullet("Pada menu Siswa, admin juga dapat menghapus semua akun siswa pada suatu kelas secara massal jika diperlukan (misalnya saat proses kelulusan siswa di akhir tahun ajaran).", "2. ")

    # --- NEW SECT 7: LUPA PASSWORD ---
    add_h1("7. Panduan Pemulihan Kata Sandi (Lupa Password) Berdasarkan Peran")
    add_body("Sistem Absenku menyediakan mekanisme pengubahan kata sandi yang disesuaikan dengan wewenang masing-masing peran pengguna demi keamanan sistem.")

    add_h2("A. Peran Siswa dan Guru / Wali Kelas (Pemulihan Mandiri via WhatsApp OTP)")
    add_body("Siswa dan Guru/Wali Kelas dapat menyetel ulang kata sandi mereka secara mandiri melalui website dengan memanfaatkan pengiriman kode OTP ke nomor WhatsApp terdaftar:")
    add_number("Pada halaman masuk (Login), ketuk tombol atau tautan Lupa Password?.", "1. ")
    add_number("Isi form verifikasi identitas dengan memasukkan Email terdaftar, Nama Lengkap Anda (sesuai profil), serta Nomor Identitas (NISN untuk siswa, atau NIP untuk guru).", "2. ")
    add_number("Sistem akan mencocokkan data. Jika data valid dan nomor WhatsApp terdaftar di sistem, sebuah pesan WhatsApp berisi 6 digit Kode OTP akan dikirimkan ke HP Anda.", "3. ")
    add_number("Masukkan 6 digit Kode OTP tersebut pada halaman verifikasi di website. Kode ini memiliki batas kedaluwarsa 30 detik untuk alasan keamanan.", "4. ")
    add_number("Jika verifikasi OTP sukses, Anda akan langsung diarahkan untuk membuat password baru yang aman. Klik Simpan untuk memperbarui kata sandi.", "5. ")
    add_note_box("Catatan Penting Pemulihan Mandiri:", "Apabila Siswa atau Guru belum pernah mengisi nomor HP/WhatsApp di profil mereka, maka kode OTP tidak dapat terkirim. Pengguna yang bersangkutan harus menghubungi Administrator Sekolah untuk mereset password ke sandi bawaan.")

    add_h2("B. Peran Petugas Piket (Melalui Administrator)")
    add_body("Demi menjaga keamanan sistem harian, akun Petugas Piket tidak memiliki fitur pemulihan kata sandi secara mandiri menggunakan OTP WhatsApp:")
    add_bullet("Petugas Piket yang lupa kata sandinya wajib melapor kepada Administrator Utama sekolah.", "1. ")
    add_bullet("Admin Utama kemudian akan mencarinya pada menu Pengguna, mengedit detail akun tersebut, lalu menuliskan password baru untuk akun Petugas Piket tersebut secara langsung.", "2. ")

    add_h2("C. Peran Administrator Sekolah / Admin Utama (Melalui Sistem Database)")
    add_body("Karena akun Admin memiliki kekuasaan penuh atas sistem dan basis data, fitur reset kata sandi mandiri di halaman login dikunci demi menghindari risiko eksploitasi:")
    add_bullet("Bila Admin Utama lupa password, tidak ada mekanisme lupa password yang dapat digunakan di website.", "1. ")
    add_bullet("Penggantian kata sandi admin wajib dilakukan secara teknis oleh Operator Server atau Tim IT Pengelola melalui database backend secara langsung (misalnya melalui database seeder atau query hashing password baru).", "2. ")

    # --- SECT 8: TROUBLESHOOTING ---
    add_h1("8. Tips & Solusi Masalah (Troubleshooting)")
    
    add_h2("Masalah A: HP Siswa Menampilkan Pesan 'Anda berada di luar radius sekolah' padahal sudah di sekolah")
    add_body("Penyebab: GPS HP siswa belum sinkron, koneksi internet lambat, atau siswa berada di dalam kelas beton yang tebal menghalangi sinyal GPS.")
    add_bullet("Mintalah siswa keluar kelas sejenak ke koridor atau lapangan terbuka agar HP mudah menangkap sinyal satelit.", "Solusi 1: ")
    add_bullet("Minta siswa membuka aplikasi Google Maps terlebih dahulu di HP mereka untuk mempercepat penentuan posisi GPS, lalu kembali ke halaman absensi Absenku.", "Solusi 2: ")
    add_bullet("Lakukan penyegaran (Refresh/Reload) pada halaman absensi di browser.", "Solusi 3: ")

    add_h2("Masalah B: Izin Lokasi di Browser HP Terlanjur Ditolak (Blocked)")
    add_body("Penyebab: Siswa atau Guru tidak sengaja memilih tombol 'Block/Deny' saat browser meminta izin lokasi.")
    add_bullet("Buka aplikasi browser Chrome/Safari di HP Anda.", "Solusi untuk Chrome (Android): ")
    add_bullet("Ketuk ikon titik tiga di kanan atas -> pilih Pengaturan (Settings) -> pilih Setelan Situs (Site Settings) -> pilih Lokasi (Location). Cari alamat web absensi sekolah dan ubah izinnya menjadi 'Izinkan' (Allow).", "Lanjutan Chrome: ", indent=2)
    add_bullet("Buka Pengaturan HP (Settings) -> pilih Safari -> pilih Location (Lokasi) -> ganti menjadi 'Ask' atau 'Allow'.", "Solusi untuk Safari (iPhone): ")

    add_h2("Masalah C: Pengguna Mengalami Kendala Akses Lainnya")
    add_bullet("Pastikan tidak ada salah penulisan karakter besar/kecil (Case Sensitive) saat mengetikkan email, NIP/NISN, atau nama lengkap Anda ketika mencoba login atau menggunakan fitur pemulihan.", "Tips 1: ")
    add_bullet("Hubungi Admin atau Wali Kelas jika akun Anda dinonaktifkan atau memerlukan penggantian informasi utama.", "Tips 2: ")

    # Save
    filepath = "/home/fauzanms/Documents/absenku/Buku_Panduan_Penggunaan_Absenku.docx"
    doc.save(filepath)
    print(f"Success! Manual book created at: {filepath}")

if __name__ == "__main__":
    generate_manual()
